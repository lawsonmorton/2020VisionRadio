import os
import smtplib
import datetime
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from google import genai
from google.genai import types

# 1. INITIALIZE GEMINI API
client = genai.Client(api_key=os.environ.get("GEMINI_API_KEY"))

# Determine dates for tomorrow's 8:00 AM broadcast
today = datetime.date.today()
tomorrow = today + datetime.timedelta(days=1)
date_str = tomorrow.strftime("%A, %B %d, %Y").upper()

# 2. CONSTRUCT SYSTEM PROMPT WITH SOURCING RULES
system_prompt = f"""
You are an expert Caribbean radio news presenter writing a broadcast script for 2020 Vision Radio.
Tomorrow's broadcast date is {date_str}. The reading broadcast time is 8:00 AM.

FORMAT RULES:
- Script formatted for a DJ/Presenter to read on-air.
- Use plain spoken English for on-air text.
- Include natural transitions and stage cues in brackets like [pause] and [transition sting].
- MUST include source URLs in brackets next to each story for the presenter to verify on-air source credibility (e.g. [Source: https://...]).

COVERAGE AREA & STRICT ORDER:
1. St. Kitts and Nevis (2-4 real stories)
2. Antigua and Barbuda (2-4 real stories)
3. Anguilla (2-3 real stories)
4. Sint Maarten / St. Martin (2-3 real stories)
5. Sint Eustatius (Statia) (1-2 real stories)
6. Saba (1-2 real stories)
7. Montserrat (1-2 real stories)
8. Wider Caribbean News (Include key stories from Dominica, Jamaica, St. Vincent and the Grenadines, etc.)
9. Weather Segment: Focus EXCLUSIVELY on St. Kitts and Nevis weather outlook for tomorrow (High/Low, Wind, Rain %, Hurricane advisories).

Provide the output formatted into two distinct sections: SCRIPT 1: NEWS AND WEATHER, and SCRIPT 2: REGIONAL SPORTS (leading with SKN Sports).
"""

# Call Gemini API with Web Search enabled for up-to-date real stories
response = client.models.generate_content(
    model="gemini-2.5-flash",
    contents=system_prompt,
    config=types.GenerateContentConfig(
        tools=[{"google_search": {}}],  # Enables real-time web verification
    ),
)
script_content = response.text


# 3. BUILD WORD DOCUMENT (.DOCX)
doc = docx.Document()

# Configure page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_header_banner(title_text, subtitle_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.5)
    cell = tbl.cell(0, 0)

    # Dark blue background banner
    shading = parse_xml(r'<w:shd {} w:fill="1B365D"/>'.format(nsdecls("w")))
    cell._tc.get_or_add_tcPr().append(shading)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p.add_run(title_text)
    run_title.font.name = "Arial"
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p2.add_run(subtitle_text)
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(0xD0, 0xE1, 0xF9)


# Title Banner
add_header_banner(
    "2020 VISION RADIO - DAILY NEWS & WEATHER SCRIPT",
    f"FOR BROADCAST: {date_str} AT 8:00 AM",
)

# Parse script text into paragraphs with highlighted cues and sources
lines = script_content.split("\n")
for line in lines:
    if line.strip():
        p = doc.add_paragraph()
        parts = line.split("[")
        p.add_run(parts[0])
        for part in parts[1:]:
            if "]" in part:
                cue_text, rest = part.split("]", 1)
                cue_full = "[" + cue_text + "]"

                r_cue = p.add_run(cue_full)
                r_cue.font.bold = True

                # Highlight cues vs source links
                if cue_text.startswith("Source:"):
                    r_cue.font.italic = True
                    r_cue.font.size = Pt(9)
                    r_cue.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                else:
                    r_cue.font.color.rgb = RGBColor(
                        0xC0, 0x39, 0x2B
                    )  # Dark Red for stage cues

                p.add_run(rest)
            else:
                p.add_run("[" + part)

doc_filename = f"2020_Vision_Radio_Script_{tomorrow.strftime('%Y_%m_%d')}.docx"
doc.save(doc_filename)


# 4. SEND EMAIL WITH ATTACHMENT
SENDER_EMAIL = os.environ.get("SENDER_EMAIL")
SENDER_PASSWORD = os.environ.get("SENDER_PASSWORD")
PRESENTER_EMAIL = os.environ.get("PRESENTER_EMAIL")

msg = MIMEMultipart()
msg["From"] = SENDER_EMAIL
msg["To"] = PRESENTER_EMAIL
msg["Subject"] = f"2020 Vision Radio Script - For {tomorrow.strftime('%A, %b %d')}"

body = f"""Hello Team,

Attached is tomorrow's automated news, sports, and weather script for 2020 Vision Radio ({date_str}).

Summary of Coverage:
- Core Islands (St. Kitts & Nevis, Antigua & Barbuda, Anguilla, Sint Maarten, Statia, Saba, Montserrat)
- Wider Caribbean News (Dominica, Jamaica, St. Vincent)
- Weather Segment: St. Kitts and Nevis Focus Only
- Branded Source Links included in brackets for source verification.

Broadcast Time: 8:00 AM tomorrow.
"""

msg.attach(MIMEText(body, "plain"))

# Attach Word Doc
with open(doc_filename, "rb") as f:
    part = MIMEBase("application", "octet-stream")
    part.set_payload(f.read())
    encoders.encode_base64(part)
    part.add_header(
        "Content-Disposition", f"attachment; filename={doc_filename}"
    )
    msg.attach(part)

# Send via SMTP (Gmail default)
server = smtplib.SMTP("smtp.gmail.com", 587)
server.starttls()
server.login(SENDER_EMAIL, SENDER_PASSWORD)
server.sendmail(SENDER_EMAIL, PRESENTER_EMAIL, msg.as_string())
server.quit()

print("Script generated and emailed successfully!")