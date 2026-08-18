import datetime
import json
import os
import re
import smtplib
import time
import traceback
from email import encoders
from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
import feedparser
import requests

# --- 1. ENVIRONMENT VARIABLES ---
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "").strip()
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "").strip()
ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "").strip()

SENDER_EMAIL = os.environ.get("SENDER_EMAIL", "").strip()
SENDER_PASSWORD = os.environ.get("SENDER_PASSWORD", "").strip()
PRESENTER_EMAIL_RAW = os.environ.get("PRESENTER_EMAIL", "")

# RSS Feed Links
SKN_POLICE_FB_RSS = os.environ.get("SKN_POLICE_FB_RSS", "").strip()
NEVIS_NEWSCAST_RSS = os.environ.get("NEVIS_NEWSCAST_RSS", "").strip()
TV_CARIBBEAN_RSS = os.environ.get("TV_CARIBBEAN_RSS", "").strip()

recipients = [
    email.strip() for email in PRESENTER_EMAIL_RAW.split(",") if email.strip()
]


def add_hyperlink(paragraph, url, text, color="0000FF", underline=True):
    """Adds a native XML hyperlink to a Word document."""
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )

    hyperlink = parse_xml(
        '<w:hyperlink'
        ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
        f' r:id="{r_id}"/>'
    )

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)

    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)

    rPr.append(
        parse_xml(r'<w:rFonts %s w:ascii="Arial" w:hAnsi="Arial"/>' % nsdecls("w"))
    )
    rPr.append(parse_xml(r'<w:sz %s w:val="18"/>' % nsdecls("w")))  # 9pt
    rPr.append(parse_xml(r'<w:i %s/>' % nsdecls("w")))

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def verify_and_clean_links(script_text):
    """
    Scans the generated script for URLs in [Source: ...] tags, tests each HTTP status,
    and replaces dead or unreachable links with '[Link Unavailable]' to guarantee working links.
    """
    url_pattern = re.compile(r"https?://[^\s\]\|]+")
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/120.0.0.0 Safari/537.36"
        )
    }

    validated_urls = {}

    def check_url(match):
        raw_url = match.group(0)
        url = raw_url.rstrip(".,;)")  # Clean trailing punctuation

        if url in validated_urls:
            return validated_urls[url]

        print(f"Validating link: {url}")
        try:
            # 1. Try HEAD request first for fast checking
            response = requests.head(
                url, timeout=4, headers=headers, allow_redirects=True
            )
            if response.status_code < 400:
                validated_urls[url] = url
                return url

            # 2. Fallback to stream GET request if HEAD is rejected by host
            response = requests.get(
                url, timeout=4, headers=headers, stream=True
            )
            if response.status_code < 400:
                validated_urls[url] = url
                return url
        except Exception as err:
            print(f"   Link check failed ({err}) for: {url}")

        print(f"--> REMOVED BROKEN LINK: {url}")
        validated_urls[url] = "[Link Unavailable]"
        return "[Link Unavailable]"

    return url_pattern.sub(check_url, script_text)


# --- 2. FETCH OFFICIAL RSS & MEDIA FEEDS ---
media_context_items = []
rss_sources = [
    ("SKN Police Force", SKN_POLICE_FB_RSS),
    ("Nevis Newscast", NEVIS_NEWSCAST_RSS),
    ("Television Caribbean", TV_CARIBBEAN_RSS),
]

for source_name, rss_url in rss_sources:
    if rss_url:
        try:
            print(f"Fetching {source_name} feed: {rss_url}")
            feed = feedparser.parse(rss_url)
            for entry in feed.entries[:5]:  # Fetch top 5 items
                title = getattr(entry, "title", "")
                summary = getattr(entry, "summary", "") or getattr(
                    entry, "description", ""
                )
                link = getattr(entry, "link", "")
                clean_summary = re.sub("<[^<]+?>", "", summary)[:300]
                media_context_items.append(
                    f"- [{source_name}] {title}: {clean_summary} (Source: {link})"
                )
        except Exception as rss_err:
            print(
                f"Notice: Could not parse {source_name} feed ({rss_url}): {rss_err}"
            )

fb_context_str = ""
if media_context_items:
    fb_context_str = (
        "\nLATEST OFFICIAL BULLETINS, POLICE REPORTS & MEDIA FEEDS:\n"
        + "\n".join(media_context_items)
        + "\n"
    )

# --- 3. LOAD STORY HISTORY ---
HISTORY_FILE = "seen_stories.json"
recent_topics = []
if os.path.exists(HISTORY_FILE):
    try:
        with open(HISTORY_FILE, "r") as f:
            recent_topics = json.load(f).get("recent_topics", [])[-20:]
    except Exception:
        pass

history_context = ""
if recent_topics:
    topics_formatted = "\n- ".join(recent_topics)
    history_context = (
        f"\nRECENTLY COVERED TOPICS (STRICT DEDUPLICATION - DO NOT REPEAT UNLESS"
        f" THERE IS A MAJOR NEW DEVELOPMENT):\n- {topics_formatted}\n"
    )

today = datetime.date.today()
tomorrow = today + datetime.timedelta(days=1)
date_str = tomorrow.strftime("%A, %B %d, %Y").upper()

# --- 4. SYSTEM PROMPT ---
system_prompt = f"""You are an expert Caribbean radio news presenter writing a broadcast script for 2020 Vision Radio 102.1 FM.
Tomorrow's broadcast date is {date_str}. Scheduled Broadcast Window starts at 8:00 AM.

STRICT REGIONAL NEWS COVERAGE QUOTAS:
1. ST. KITTS AND NEVIS (AT LEAST 6 ITEMS TOTAL):
   - You MUST include at least 3 distinct news stories for St. Kitts.
   - You MUST include at least 3 distinct news stories for Nevis.
2. LEEWARD & NEIGHBORING ISLANDS: Include stories for Antigua & Barbuda, Anguilla, Sint Maarten/St. Martin, Statia, Saba, and Montserrat.
3. MANDATORY WIDER CARIBBEAN COVERAGE (AT LEAST 1 ITEM EACH):
   You MUST attempt to include at least 1 real news item for each of the following territories:
   - Trinidad and Tobago
   - St. Vincent and the Grenadines
   - Dominica
   - Jamaica
   - Grenada
   - St. Lucia
   - Guyana
   - Belize
   - United States Virgin Islands (USVI)
   - British Virgin Islands (BVI)

STRICT WEATHER RULES:
- Focus EXCLUSIVELY on St. Kitts and Nevis.
- DAILY WEATHER RULE: Provide the forecast strictly for tomorrow's date ({date_str}) ONLY. Do NOT include a 5-day forecast or general multi-day outlook.
- THURSDAY SPECIAL RULE: If tomorrow's date is Friday (script generated on Thursday), you MUST create a dedicated section immediately after Friday's forecast titled 'Weekend Outlook', providing a detailed weather forecast specifically for Saturday and Sunday.
- Forecast details must include synoptic analysis, island-by-island breakdown (Basseterre, Sandy Point, Nevis Peak, Charlestown), marine swell advisories, and tide schedules.

4. SPORTS SEGMENT: Lead with St. Kitts & Nevis local athletics and community cricket/football leagues, followed by Leeward Islands Cricket, CWI, CONCACAF, and regional sports updates.

ACCURACY & FORMATTING RULES:
- STRICT DEDUPLICATION: Do NOT repeat news items covered in previous days unless there is a fresh, breaking update.
- DO NOT mention segment durations, timing lengths, or word counts inside the generated script text (e.g. do NOT write '30 Minutes', '7 Minutes', or '4000 words').
- DO NOT invent fictional persons or names (e.g., do NOT use hallucinated names like 'Marcus St. Clair'). Report ONLY real, verifiable public officials and real citizens.
- Write in clear, spoken radio-presenter English for 2020 Vision Radio 102.1 FM.
- Include stage cues in brackets like [pause], [station intro sting], [transition sting], and commercial break placeholders.
- LINK ACCURACY MANDATE: Include direct source URLs ONLY if retrieved directly from live RSS context feeds or verified live search results. NEVER fabricate, guess, or synthesize URL strings.
- MEDIA LINKS REQUIREMENT: If an item includes verified associated audio or video, include direct access links inside the source brackets (e.g. [Source: https://... | Audio: https://...]).
{fb_context_str}{history_context}
BROADCAST STRUCTURE:
=== SCRIPT 1: NEWS SEGMENT ===
=== SCRIPT 2: WEATHER SEGMENT ===
=== SCRIPT 3: SPORTS SEGMENT ===

At the very end of your response, output a single JSON block:
```json
{{"topics": ["Topic 1", "Topic 2"]}}
```"""

# --- 5. EXECUTE AI GENERATION WITH FALLBACKS ---
script_content = None

# ATTEMPT 1: GOOGLE GEMINI
if GEMINI_API_KEY:
    try:
        print("Attempting Primary Provider: Google Gemini...")
        from google import genai
        from google.genai import types

        client = genai.Client(api_key=GEMINI_API_KEY)

        available_models = []
        try:
            for m in client.models.list():
                m_name = getattr(m, "name", "")
                clean_name = m_name.replace("models/", "")
                if "flash" in clean_name.lower() or "pro" in clean_name.lower():
                    available_models.append(clean_name)
        except Exception as list_err:
            print(f"Could not list models dynamically: {list_err}")

        if not available_models:
            available_models = [
                "gemini-2.5-flash",
                "gemini-2.5-flash-lite",
                "gemini-2.0-flash-exp",
            ]

        print(f"Discovered active Gemini candidate models: {available_models}")

        for g_model in available_models:
            try:
                response = client.models.generate_content(
                    model=g_model,
                    contents=system_prompt,
                    config=types.GenerateContentConfig(tools=[{"google_search": {}}]),
                )
                if response and hasattr(response, "text") and response.text:
                    script_content = response.text
                    print(f"SUCCESS: Generated via Google Gemini ({g_model})!")
                    break
            except Exception as g_err1:
                try:
                    response = client.models.generate_content(
                        model=g_model, contents=system_prompt
                    )
                    if response and hasattr(response, "text") and response.text:
                        script_content = response.text
                        print(
                            f"SUCCESS: Generated via Google Gemini ({g_model} - plain)!"
                        )
                        break
                except Exception as g_err2:
                    print(f"Gemini model {g_model} failed: {g_err1}")

    except Exception as e:
        print(f"Gemini API initialization failed: {e}")

# ATTEMPT 2: OPENAI CHATGPT (FALLBACK)
if not script_content and OPENAI_API_KEY:
    try:
        print("Switching to Fallback Provider 1: OpenAI ChatGPT...")
        from openai import OpenAI

        client = OpenAI(api_key=OPENAI_API_KEY, timeout=30.0, max_retries=2)
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": system_prompt}],
        )
        if (
            response
            and response.choices
            and response.choices[0].message
            and response.choices[0].message.content
        ):
            script_content = response.choices[0].message.content
            print("SUCCESS: Generated via OpenAI ChatGPT!")
    except Exception as e:
        print(f"OpenAI API failed: {e}")

# ATTEMPT 3: ANTHROPIC CLAUDE (FALLBACK)
if not script_content and ANTHROPIC_API_KEY:
    try:
        print("Switching to Fallback Provider 2: Anthropic Claude...")
        import anthropic

        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        anthropic_models = [
            "claude-3-5-sonnet-latest",
            "claude-3-5-haiku-latest",
            "claude-3-haiku-20240307",
        ]
        for model_name in anthropic_models:
            try:
                response = client.messages.create(
                    model=model_name,
                    max_tokens=4000,
                    messages=[{"role": "user", "content": system_prompt}],
                )
                if response and response.content and response.content[0].text:
                    script_content = response.content[0].text
                    print(f"SUCCESS: Generated via Anthropic Claude ({model_name})!")
                    break
            except Exception as inner_e:
                print(f"Anthropic model {model_name} failed: {inner_e}")
    except Exception as e:
        print(f"Anthropic API failed: {e}")

if not script_content:
    print("CRITICAL ERROR: All AI providers failed or credentials missing.")
    exit(1)

# Extract history topics
clean_script = script_content
if "```json" in script_content:
    try:
        parts = script_content.split("```json")
        clean_script = parts[0].strip()
        json_str = parts[1].split("```")[0].strip()
        extracted_data = json.loads(json_str)
        recent_topics.extend(extracted_data.get("topics", []))
        with open(HISTORY_FILE, "w") as f:
            json.dump({"recent_topics": recent_topics[-25:]}, f, indent=2)
    except Exception:
        pass

# --- 5.5 VERIFY AND CLEAN LINKS ---
print("Verifying all generated source URLs...")
clean_script = verify_and_clean_links(clean_script)

# --- 6. BUILD WORD DOCUMENT (.DOCX) ---
doc = docx.Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

tbl = doc.add_table(rows=1, cols=1)
tbl.autofit = False
tbl.columns[0].width = Inches(6.5)
cell = tbl.cell(0, 0)
shading = parse_xml(r'<w:shd {} w:fill="1B365D"/>'.format(nsdecls("w")))
cell._tc.get_or_add_tcPr().append(shading)

p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p.add_run("2020 VISION RADIO 102.1 FM - DAILY BROADCAST SCRIPT")
r1.font.name = "Arial"
r1.font.size = Pt(16)
r1.font.bold = True
r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run(f"FOR BROADCAST: {date_str} (News | Weather | Sports)")
r2.font.name = "Arial"
r2.font.size = Pt(11)
r2.font.bold = True
r2.font.color.rgb = RGBColor(0xD0, 0xE1, 0xF9)

url_pattern = re.compile(r"https?://[^\s\]\|]+")

for line in clean_script.split("\n"):
    if line.strip():
        p_line = doc.add_paragraph()
        parts = line.split("[")
        p_line.add_run(parts[0])
        for part in parts[1:]:
            if "]" in part:
                cue_text, rest = part.split("]", 1)
                if cue_text.startswith("Source:"):
                    r_bracket = p_line.add_run("[")
                    r_bracket.font.italic = True
                    r_bracket.font.bold = True
                    r_bracket.font.size = Pt(9)
                    r_bracket.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                    last_idx = 0
                    for match in url_pattern.finditer(cue_text):
                        start, end = match.span()
                        if start > last_idx:
                            r_prefix = p_line.add_run(cue_text[last_idx:start])
                            r_prefix.font.italic = True
                            r_prefix.font.bold = True
                            r_prefix.font.size = Pt(9)
                            r_prefix.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                        found_url = match.group(0)
                        add_hyperlink(p_line, found_url, found_url)
                        last_idx = end

                    if last_idx < len(cue_text):
                        r_suffix = p_line.add_run(cue_text[last_idx:])
                        r_suffix.font.italic = True
                        r_suffix.font.bold = True
                        r_suffix.font.size = Pt(9)
                        r_suffix.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                    r_close = p_line.add_run("]")
                    r_close.font.italic = True
                    r_close.font.bold = True
                    r_close.font.size = Pt(9)
                    r_close.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                else:
                    r_cue = p_line.add_run(f"[{cue_text}]")
                    r_cue.font.bold = True
                p_line.add_run(rest)
            else:
                p_line.add_run(f"[{part}")

output_filename = f"Broadcast_Script_{tomorrow.strftime('%Y%m%d')}.docx"
doc.save(output_filename)
print(f"Successfully generated script: {output_filename}")

# --- 7. SEND EMAIL WITH ATTACHMENT ---
if SENDER_EMAIL and SENDER_PASSWORD and recipients:
    print(f"Sending email to: {', '.join(recipients)}...")
    try:
        msg = MIMEMultipart()
        msg["From"] = SENDER_EMAIL
        msg["To"] = ", ".join(recipients)
        msg["Subject"] = f"2020 Vision Radio Script - {date_str}"

        body_text = (
            f"Hello,\n\n"
            f"Please find attached the daily broadcast script for {date_str}.\n\n"
            f"Best regards,\n2020 Vision Radio Automation"
        )
        msg.attach(MIMEText(body_text, "plain"))

        if os.path.exists(output_filename):
            with open(output_filename, "rb") as attachment:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header(
                "Content-Disposition",
                f"attachment; filename={os.path.basename(output_filename)}",
            )
            msg.attach(part)

        server = smtplib.SMTP("smtp.gmail.com", 587)
        server.starttls()
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.sendmail(SENDER_EMAIL, recipients, msg.as_string())
        server.quit()
        print("Email sent successfully!")

    except Exception as email_err:
        print(f"ERROR: Failed to send email: {email_err}")
        traceback.print_exc()
else:
    print(
        "Notice: Email sending skipped. Please check SENDER_EMAIL, SENDER_PASSWORD, and PRESENTER_EMAIL secrets in GitHub repository."
    )